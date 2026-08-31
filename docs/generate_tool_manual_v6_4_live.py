#!/usr/bin/env python3
from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageEnhance, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF

ROOT = Path(__file__).resolve().parent
IMG_DIR = ROOT / "images" / "manual_v6_4_live"
META_FILE = IMG_DIR / "annotation_map.json"
ANN_DIR = ROOT / "images" / "manual_v6_4_live_annotated"
VIEW_DIR = ROOT / "images" / "manual_v6_4_live_view"
DOCX_OUT = ROOT / "Coverity_Tool_User_Manual.docx"
PDF_OUT = ROOT / "Coverity_Tool_User_Manual.pdf"

IMG_ORDER = [
    "01_setup_live.png",
    "02_commit_live.png",
    "03a_pull_top_live.png",
    "03b_pull_bottom_live.png",
    "04_analysis_live.png",
    "05_results_live.png",
    "06_detail_live.png",
    "07a_push_csv_top_live.png",
    "07b_push_csv_bottom_live.png",
    "08a_direct_push_top_live.png",
    "08b_direct_push_bottom_live.png",
]

IMG_TITLES = {
    "01_setup_live.png": "Setup Page",
    "02_commit_live.png": "Commit Defects Dialog",
    "03a_pull_top_live.png": "Pull Defects Dialog - Connection and Output Setup",
    "03b_pull_bottom_live.png": "Pull Defects Dialog - Pull Action and Log",
    "04_analysis_live.png": "Analysis Progress",
    "05_results_live.png": "Disposition Results",
    "06_detail_live.png": "Full Detail Window",
    "07a_push_csv_top_live.png": "Push Dispositions (CSV) - Connection and CSV Review",
    "07b_push_csv_bottom_live.png": "Push Dispositions (CSV) - Validation and Push Controls",
    "08a_direct_push_top_live.png": "Direct Push Results - Connection and Push Mode",
    "08b_direct_push_bottom_live.png": "Direct Push Results - Validation and Push Controls",
}

GUIDANCE = {
    "01_setup_live.png": [
        "Use this screen to define the full run context: report, source root, and output folder.",
        "If findings are not yet in Coverity Connect, use Commit Defects first, then return to Pull Defects.",
        "Source Code Root should match the code revision that produced the report.",
        "Start Disposition only after all required fields are populated.",
    ],
    "02_commit_live.png": [
        "Commit is used only when you already have a Coverity intermediate directory (idir).",
        "Connect first to load valid project and stream values from server.",
        "Resolve all validation warnings before clicking Commit to Coverity.",
        "Use the output pane to troubleshoot tool-path, auth, and stream issues.",
    ],
    "03a_pull_top_live.png": [
        "Pull fetches defects directly from Coverity Connect and creates a structured Excel input.",
        "Run Test Connection before selecting project and stream.",
        "Set a practical defect limit for initial pull and increase later if needed.",
        "Keep the generated pull Excel path; this becomes the Setup input.",
    ],
    "03b_pull_bottom_live.png": [
        "The lower part of the Pull dialog contains the action button, progress bar, and run log.",
        "Use this area to confirm the pull started and completed successfully.",
        "If the pull fails, copy the visible log message and check server, credential, or permission setup.",
    ],
    "04_analysis_live.png": [
        "Progress area shows processed defects, percentage, elapsed time, and ETC.",
        "Log panel reports file loading, indexing, and per-defect classification traces.",
        "Warnings in red or yellow should be reviewed before final push.",
        "Cancel should be used only when wrong input/source path is selected.",
    ],
    "05_results_live.png": [
        "Left panel is your CID work queue; select one finding at a time.",
        "Top filters reduce review scope by classification and category.",
        "Center panel contains disposition rationale, proposed fix, and actions.",
        "Right panel verifies exact code context before acceptance or override.",
    ],
    "06_detail_live.png": [
        "Use this window for high-confidence review of one CID with full source context.",
        "Read Analysis and Source-validated Proposed Fix before changing disposition.",
        "Accept Suggestion records the current recommendation as final.",
        "Override is for reviewer-driven corrections with explicit rationale.",
    ],
    "07a_push_csv_top_live.png": [
        "Use this dialog when pushing decisions from coverity_final_decisions.csv.",
        "Connect and select correct project, stream, and triage store first.",
        "Validate CIDs against server to detect stale or wrong-project entries.",
        "Push only after validation is clean and table rows are reviewed.",
    ],
    "07b_push_csv_bottom_live.png": [
        "The lower CSV Push view shows the defect rows and final push buttons.",
        "Review the validation status and table rows before uploading dispositions.",
        "Use Dry Run when you need a pre-flight check without changing Coverity triage state.",
    ],
    "08a_direct_push_top_live.png": [
        "Use direct push to submit in-memory review results without CSV export.",
        "Choose push mode based on policy: accepted only, decided only, or all analysed defects.",
        "Run CID validation before push so unmatched CIDs are excluded.",
        "Dry run is useful for pre-flight checks in controlled environments.",
    ],
    "08b_direct_push_bottom_live.png": [
        "The lower Direct Push view shows matched rows, status, and the final push controls.",
        "Only validated server CIDs should be pushed.",
        "After push, review the result message and any failed rows before closing the dialog.",
    ],
}

WORKFLOW = {
    "01_setup_live.png": [
        ("Select input", "Browse to the pulled Excel file or HTML report folder.", "The selected path appears in the Coverity Report field."),
        ("Select source root", "Browse to the matching source checkout.", "The tool can show source context during review."),
        ("Select output folder", "Choose a writable folder for generated CSV and logs.", "All analysis outputs are saved in one known location."),
        ("Start analysis", "Click Start Disposition after fields are complete.", "The analysis progress screen opens."),
    ],
    "02_commit_live.png": [
        ("Set Coverity tools", "Point to Coverity bin only when commands are not already on PATH.", "The tool can find cov-commit-defects."),
        ("Choose idir", "Select the intermediate directory created by Coverity analysis.", "The commit source is ready for upload."),
        ("Connect", "Enter server, port, and credentials, then connect.", "Project and stream selections become available."),
        ("Commit", "Select project and stream, review readiness, then commit.", "Defects are uploaded to Coverity Connect."),
    ],
    "03a_pull_top_live.png": [
        ("Connect", "Enter host, port, username, and password, then test connection.", "Project and stream lists load from server."),
        ("Select scope", "Choose project, stream, and defect limit.", "The pull is restricted to the intended Coverity data set."),
        ("Choose output", "Confirm the Excel save path.", "The pulled report location is known before download."),
    ],
    "03b_pull_bottom_live.png": [
        ("Start pull", "Click Pull Defects after connection and stream selection are valid.", "Download begins and progress updates."),
        ("Watch log", "Read the black log area for server or permission messages.", "Successful completion confirms the Excel file is ready."),
        ("Close dialog", "Close after the generated file is available.", "Return to Setup and use the Excel file as input."),
    ],
    "04_analysis_live.png": [
        ("Monitor progress", "Watch processed count, percentage, elapsed time, and ETC.", "You know whether the run is moving normally."),
        ("Review messages", "Read the analysis log while the run proceeds.", "Input, source, and classification warnings are visible."),
        ("Let it finish", "Avoid cancelling unless the wrong input was selected.", "Results open after analysis completes."),
    ],
    "05_results_live.png": [
        ("Filter work queue", "Use status/category filters to focus review.", "Only relevant CIDs remain visible."),
        ("Review one CID", "Select a finding and read disposition, rationale, and source panel.", "Reviewer understands why the tool suggested that decision."),
        ("Accept or override", "Accept valid suggestions or override with a reviewer decision.", "Final decision is captured for export or push."),
        ("Open details", "Use the full detail window for complex findings.", "More source and evidence are available."),
    ],
    "06_detail_live.png": [
        ("Confirm identity", "Check CID, checker, and classification in the header.", "You are reviewing the intended finding."),
        ("Read evidence", "Review analysis text and source-validated proposed fix.", "Decision is based on recorded evidence."),
        ("Check source", "Compare the highlighted source context with the finding.", "Wrong-source or stale-code issues are detected."),
        ("Decide", "Accept suggestion or override with reviewer judgement.", "The decision returns to the results workflow."),
    ],
    "07a_push_csv_top_live.png": [
        ("Connect", "Enter server details and connect.", "Project, stream, and triage store can be selected."),
        ("Load CSV", "Select coverity_final_decisions.csv from the output folder.", "Rows appear in the push table."),
        ("Validate CIDs", "Run validation before push.", "Invalid or mismatched CIDs are detected before upload."),
    ],
    "07b_push_csv_bottom_live.png": [
        ("Review rows", "Check CID, server CID, classification, action, comment, checker, and file.", "Only intended dispositions are queued."),
        ("Dry run", "Use Dry Run when required by team process.", "The tool checks what would be pushed."),
        ("Push", "Push only after validation and row review are complete.", "Final decisions are submitted to Coverity."),
    ],
    "08a_direct_push_top_live.png": [
        ("Connect", "Enter server credentials and connect.", "Coverity project and triage store are available."),
        ("Select mode", "Choose accepted only, decided only, or all analysed defects.", "The push set matches review policy."),
        ("Validate", "Validate CIDs against the server before upload.", "Only server-matched CIDs are ready."),
    ],
    "08b_direct_push_bottom_live.png": [
        ("Review selected rows", "Confirm each mapped CID and action before upload.", "Wrong findings are caught before server update."),
        ("Dry run", "Run a dry push first when you need confirmation.", "No Coverity data is changed."),
        ("Push direct", "Submit current in-memory decisions.", "Results are pushed without using a CSV file."),
    ],
}


def _set_layout(doc: Document) -> None:
    for sec in doc.sections:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
        sec.top_margin = Inches(0.4)
        sec.bottom_margin = Inches(0.4)
        sec.left_margin = Inches(0.45)
        sec.right_margin = Inches(0.45)


def _set_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)


def _title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Coverity Tool User Manual")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = p.add_run(f"Version 6.6 | Updated {date.today().isoformat()} | Detailed workflow instructions with tool screenshots")
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)


def _load_specs() -> list[dict]:
    if not META_FILE.is_file():
        raise FileNotFoundError(f"Annotation map not found: {META_FILE}")
    raw = json.loads(META_FILE.read_text(encoding="utf-8"))

    specs: list[dict] = []
    for name in IMG_ORDER:
        src_img = IMG_DIR / name
        if not src_img.is_file():
            raise FileNotFoundError(f"Missing screenshot: {src_img}")
        iw, ih = Image.open(src_img).size

        note_rows = raw.get(name)
        if not isinstance(note_rows, list) or not note_rows:
            raise ValueError(f"Missing annotation rows for {name}")

        notes: list[tuple[str, str, tuple[int, int, int, int]]] = []
        for row in note_rows:
            try:
                nm = str(row["name"])
                action = str(row["action"])
                box = tuple(int(v) for v in row["box"])
                if len(box) == 4:
                    x1, y1, x2, y2 = box
                    bw = x2 - x1
                    bh = y2 - y1
                    # Ignore clipped slivers that appear when a widget is outside viewport.
                    if bw < 20 or bh < 20:
                        continue
                    if (y2 >= ih - 1 and bh < 42) or (x2 >= iw - 1 and bw < 42):
                        continue
                    notes.append((nm, action, box))
            except Exception:
                continue
        if not notes:
            raise ValueError(f"No valid annotation boxes for {name}")

        specs.append({
            "file": name,
            "title": IMG_TITLES.get(name, name),
            "notes": notes,
            "guidance": GUIDANCE.get(name, []),
        })
    return specs


def _load_font(size: int) -> ImageFont.ImageFont:
    for face in ["arial.ttf", "segoeui.ttf", "calibri.ttf"]:
        try:
            return ImageFont.truetype(face, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_annotations(src: Path, dst: Path, notes: list[tuple[str, str, tuple[int, int, int, int]]]) -> None:
    base = Image.open(src).convert("RGB")
    # Slight contrast and sharpness boost to keep labels and fields readable in PDF.
    base = ImageEnhance.Contrast(base).enhance(1.05)
    base = ImageEnhance.Sharpness(base).enhance(1.12)

    legend_w = 430
    img = Image.new("RGB", (base.width + legend_w, base.height), "white")
    img.paste(base, (0, 0))

    draw = ImageDraw.Draw(img)
    font = _load_font(18)
    legend_font = _load_font(17)
    legend_head_font = _load_font(20)
    colors = [
        (37, 99, 235),
        (220, 38, 38),
        (22, 163, 74),
        (217, 119, 6),
        (8, 145, 178),
        (147, 51, 234),
    ]

    for idx, (_name, _action, box) in enumerate(notes, start=1):
        color = colors[(idx - 1) % len(colors)]
        x1, y1, x2, y2 = box
        draw.rectangle((x1, y1, x2, y2), outline=color, width=3)

    lx = base.width + 24
    draw.text((lx, 24), "Screen Guide", fill=(30, 58, 95), font=legend_head_font)
    y = 62
    for idx, (name, action, _box) in enumerate(notes, start=1):
        color = colors[(idx - 1) % len(colors)]
        draw.ellipse((lx, y + 2, lx + 26, y + 28), fill="white", outline=color, width=3)
        draw.text((lx + 8, y + 3), str(idx), fill=color, font=font)
        draw.text((lx + 38, y), name, fill=(15, 23, 42), font=legend_font)

        words = action.split()
        lines = []
        line = ""
        for word in words:
            trial = (line + " " + word).strip()
            if len(trial) > 38:
                lines.append(line)
                line = word
            else:
                line = trial
        if line:
            lines.append(line)
        for text_line in lines[:3]:
            y += 22
            draw.text((lx + 38, y), text_line, fill=(71, 85, 105), font=legend_font)
        y += 30

    dst.parent.mkdir(parents=True, exist_ok=True)
    img.save(dst, format="PNG", optimize=True)


def _make_clear_view(src: Path, dst_view: Path) -> None:
    img = Image.open(src).convert("RGB")
    img = ImageEnhance.Contrast(img).enhance(1.04)
    img = ImageEnhance.Sharpness(img).enhance(1.12)
    # Keep full-frame screenshots in the user manual so scroll position and layout are clear.
    view = img

    vw, vh = view.size
    max_w = 2200
    if vw > max_w:
        scale = max_w / float(vw)
        view = view.resize((int(vw * scale), int(vh * scale)), Image.Resampling.LANCZOS)

    dst_view.parent.mkdir(parents=True, exist_ok=True)
    view.save(dst_view, format="PNG", optimize=True)


def _add_overview(doc: Document) -> None:
    _add_heading(doc, "1. Data Required Before You Start")
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "Item"
    t.rows[0].cells[1].text = "Typical value"
    t.rows[0].cells[2].text = "Where used"

    rows = [
        ("Coverity host", "coverity-er.honaero.com", "Commit, Pull, Push dialogs"),
        ("Port", "443", "Commit, Pull, Push dialogs"),
        ("Credentials", "Corporate username/password or auth key", "Connection and upload actions"),
        ("Project", "Program project in Coverity", "Pull and push scoping"),
        ("Stream", "Target stream in project", "Commit target and pull scope"),
        ("Source root", "Local source tree for the same baseline", "Setup and results verification"),
        ("Output folder", "Local writable path", "Generated CSV and logs"),
    ]
    for item, value, where in rows:
        c = t.add_row().cells
        c[0].text = item
        c[1].text = value
        c[2].text = where
        for cell in c:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    _add_heading(doc, "2. End-to-End Workflow")
    for line in [
        "1) If defects are not in Coverity Connect, use Commit Defects dialog first.",
        "2) Pull defects from server to create structured Excel input.",
        "3) Run analysis and review each finding in Results and Detail windows.",
        "4) Push final decisions using CSV Push dialog or Direct Push dialog.",
    ]:
        doc.add_paragraph(line)


def _add_screen(doc: Document, idx: int, spec: dict, clear_img: Path) -> None:
    _add_heading(doc, f"{idx}. {spec['title']}")

    doc.add_picture(str(clear_img), width=Inches(10.6))
    cap = doc.add_paragraph(f"Figure {idx - 2}: {clear_img.name}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)

    if spec.get("guidance"):
        doc.add_paragraph("Purpose and review guidance:")
        for step in spec["guidance"]:
            doc.add_paragraph(step, style="List Bullet")

    steps = WORKFLOW.get(spec["file"], [])
    if steps:
        doc.add_paragraph("How to use this screen:")
        t_steps = doc.add_table(rows=1, cols=4)
        t_steps.style = "Light Grid Accent 1"
        t_steps.rows[0].cells[0].text = "Step"
        t_steps.rows[0].cells[1].text = "User action"
        t_steps.rows[0].cells[2].text = "What to enter or check"
        t_steps.rows[0].cells[3].text = "Expected result"
        for i, (action, check, result) in enumerate(steps, start=1):
            c = t_steps.add_row().cells
            c[0].text = str(i)
            c[1].text = action
            c[2].text = check
            c[3].text = result

def _add_outputs(doc: Document) -> None:
    _add_heading(doc, "11. Files Produced by the Tool")
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "File"
    t.rows[0].cells[1].text = "Content"
    t.rows[0].cells[2].text = "Use"

    rows = [
        ("coverity_pull_<stream>_<timestamp>.xlsx", "Pulled findings from Coverity", "Input for analysis"),
        ("coverity_dispositions.csv", "Initial classification suggestions", "Review and QA"),
        ("coverity_final_decisions.csv", "Final reviewer decisions", "Push source"),
        ("audit.jsonl", "Decision evidence log", "Traceability and audit"),
        ("*_pull_log.txt", "Pull diagnostics", "Connection and fetch troubleshooting"),
    ]
    for f, content, use in rows:
        c = t.add_row().cells
        c[0].text = f
        c[1].text = content
        c[2].text = use


def _add_troubleshooting(doc: Document) -> None:
    _add_heading(doc, "12. Troubleshooting")
    for line in [
        "Application does not open: keep CoverityTool.exe and _internal folder together.",
        "Cannot connect: verify host, port, credentials, and network/VPN path.",
        "No streams visible: user account may lack stream permissions in selected project.",
        "No local code in Results: verify Source Root points to matching code baseline.",
        "Push failures: validate CIDs, then check triage store and stream selection.",
    ]:
        doc.add_paragraph(line, style="List Bullet")


def _safe_pdf_text(text: str) -> str:
    return text.replace("–", "-").replace("—", "-").replace("’", "'").replace("“", '"').replace("”", '"')


def _pdf_multiline(pdf: FPDF, text: str, size: int = 9, style: str = "") -> None:
    pdf.set_font("Arial", style, size)
    pdf.set_x(10)
    pdf.multi_cell(0, 5, _safe_pdf_text(text))


def _add_pdf_table(pdf: FPDF, headers: list[str], rows: list[tuple[str, ...]], widths: list[int]) -> None:
    pdf.set_font("Arial", "B", 8)
    pdf.set_fill_color(232, 240, 254)
    pdf.set_x(10)
    for header, width in zip(headers, widths):
        pdf.cell(width, 6, _safe_pdf_text(header), border=1, fill=True)
    pdf.ln()
    pdf.set_font("Arial", "", 8)
    for row in rows:
        y0 = pdf.get_y()
        x0 = 10
        heights = []
        for value, width in zip(row, widths):
            lines = max(1, len(_safe_pdf_text(value)) // max(12, int(width / 2)) + 1)
            heights.append(lines * 4)
        row_h = max(6, min(18, max(heights)))
        for value, width in zip(row, widths):
            pdf.set_xy(x0, y0)
            pdf.multi_cell(width, 4, _safe_pdf_text(value), border=1)
            x0 += width
        pdf.set_y(y0 + row_h)
        pdf.set_x(10)


def _build_pdf(rendered: list[tuple[dict, Path]]) -> None:
    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=10)

    pdf.add_page()
    pdf.set_font("Arial", "B", 20)
    pdf.cell(0, 10, "Coverity Tool User Manual", ln=True, align="C")
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 7, f"Version 6.6 | Updated {date.today().isoformat()}", ln=True, align="C")
    pdf.ln(5)
    _pdf_multiline(pdf, "This guide explains the data needed, the screens used by the workflow, and the expected result from each user action.", 10)

    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 8, "Data Required Before You Start", ln=True)
    _add_pdf_table(pdf, ["Item", "Typical value", "Where used"], [
        ("Coverity host", "coverity-er.honaero.com", "Commit, Pull, Push dialogs"),
        ("Port", "443", "Commit, Pull, Push dialogs"),
        ("Credentials", "Corporate username/password or auth key", "Connection and upload actions"),
        ("Project", "Program project in Coverity", "Pull and push scoping"),
        ("Stream", "Target stream in project", "Commit target and pull scope"),
        ("Source root", "Local source tree for the same baseline", "Setup and results verification"),
        ("Output folder", "Local writable path", "Generated CSV and logs"),
    ], [52, 90, 120])

    for idx, (spec, view) in enumerate(rendered, start=3):
        pdf.add_page()
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 8, _safe_pdf_text(f"{idx}. {spec['title']}"), ln=True)
        pdf.image(str(view), x=8, y=20, w=280)
        pdf.set_y(160)
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 6, "Purpose and review guidance:", ln=True)
        for line in spec.get("guidance", []):
            _pdf_multiline(pdf, "- " + line, 8)
        rows = [(str(i), action, check, result) for i, (action, check, result) in enumerate(WORKFLOW.get(spec["file"], []), start=1)]
        if rows:
            pdf.ln(2)
            _add_pdf_table(pdf, ["Step", "User action", "What to enter or check", "Expected result"], rows, [14, 48, 112, 88])

    pdf.add_page()
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 8, "Files Produced by the Tool", ln=True)
    _add_pdf_table(pdf, ["File", "Content", "Use"], [
        ("coverity_pull_<stream>_<timestamp>.xlsx", "Pulled findings from Coverity", "Input for analysis"),
        ("coverity_dispositions.csv", "Initial classification suggestions", "Review and QA"),
        ("coverity_final_decisions.csv", "Final reviewer decisions", "Push source"),
        ("audit.jsonl", "Decision evidence log", "Traceability and audit"),
        ("*_pull_log.txt", "Pull diagnostics", "Connection and fetch troubleshooting"),
    ], [88, 88, 86])
    pdf.ln(6)
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 8, "Troubleshooting", ln=True)
    for line in [
        "Application does not open: keep CoverityTool.exe and _internal folder together.",
        "Cannot connect: verify host, port, credentials, and network/VPN path.",
        "No streams visible: user account may lack stream permissions in selected project.",
        "No local code in Results: verify Source Root points to matching code baseline.",
        "Push failures: validate CIDs, then check triage store and stream selection.",
    ]:
        _pdf_multiline(pdf, "- " + line, 9)

    pdf.output(str(PDF_OUT))


def main() -> None:
    specs = _load_specs()
    ANN_DIR.mkdir(parents=True, exist_ok=True)
    VIEW_DIR.mkdir(parents=True, exist_ok=True)

    rendered: list[tuple[dict, Path]] = []
    for spec in specs:
        src = IMG_DIR / spec["file"]
        if not src.is_file():
            raise FileNotFoundError(f"Missing screenshot: {src}")
        ann = ANN_DIR / spec["file"]
        view = VIEW_DIR / spec["file"]
        _draw_annotations(src, ann, spec["notes"])
        _make_clear_view(ann, view)
        rendered.append((spec, view))

    doc = Document()
    _set_style(doc)
    _set_layout(doc)
    _title(doc)
    _add_overview(doc)

    section_no = 3
    for spec, view in rendered:
        doc.add_page_break()
        _add_screen(doc, section_no, spec, view)
        section_no += 1

    _add_outputs(doc)
    _add_troubleshooting(doc)

    doc.save(str(DOCX_OUT))
    print(f"Generated docx: {DOCX_OUT}")
    print(f"Annotation map source: {META_FILE}")
    print(f"Annotated images: {ANN_DIR}")
    print(f"Clear-view images: {VIEW_DIR}")


if __name__ == "__main__":
    main()
