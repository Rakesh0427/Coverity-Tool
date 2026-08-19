#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.dml.color import ColorFormat

OUT = "/home/user/Coverity-Tool/docs/Coverity_Tool_User_Guide.docx"
IMAGES_DIR = "/home/user/Coverity-Tool/docs/images"

# Colors matching GUI palette
C_ACCENT = RGBColor(0x25, 0x63, 0xEB)  # #2563EB
C_HDR = RGBColor(0x1E, 0x3A, 0x5F) # #1E3A5F
C_BUG = RGBColor(0xDC, 0x26, 0x26)
C_FP = RGBColor(0x16, 0xA3, 0x4A)
C_INTENT = RGBColor(0xD9, 0x77, 0x06)
C_REVIEW = RGBColor(0x08, 0x91, 0xB2)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top','left','bottom','right','insideH','insideV'):
        ed = OxmlElement(f'w:{edge}')
        for k,v in kwargs.items():
            ed.set(qn(f'w:{k}'), v)
        tcPr.append(ed)

def add_horizontal_line(paragraph, color="2563EB", size="6"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_doc():
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.05

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

    # --- COVER PAGE ---
    # Top banner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COVERITY FINDINGS ANALYZER")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = C_HDR
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Automated Defect Triage for Coverity & Coverity Connect")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.italic = True

    add_horizontal_line(doc.add_paragraph(), color="2563EB", size="12")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Detailed User Guide  •  Version 1.4  •  Expert CWE Edition")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = C_ACCENT

    # Version box
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    cells = table.rows[0].cells
    cells[0].text = "Version:\n1.4 (Expert CWE)"
    cells[1].text = "Date:\n19-Aug-2026"
    cells[2].text = "Audience:\nEngineering Teams"
    for cell in cells:
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.bold = True
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'EFF6FF')
        cell._tc.get_or_add_tcPr().append(shading)
        set_cell_border(cell, val='single', sz='4', color='2563EB', space='0')
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(2.2)
    table.columns[2].width = Inches(2.2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Local • Privacy-Preserving • No Source Upload • Tree-sitter + Z3 + CWE/CERT/OWASP")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x64,0x74,0x8B)
    run.italic = True

    # What’s inside box
    p = doc.add_heading("What’s Inside This Guide", level=2)
    p.runs[0].font.color.rgb = C_HDR
    add_horizontal_line(doc.add_paragraph(), color="CBD5E1", size="4")
    bullets = [
        "How to install & run (EXE one-click vs Python venv) — p.3",
        "What counts as valid INPUT (HTML folder / Excel / Server Pull) — p.4-5",
        "Where OUTPUT goes & exact CSV/JSON formats — p.6",
        "Click-by-click GUI tour with highlighted screenshots — p.7-14",
        "How to read Bug / False positive / Intentional / Needs review — p.15",
        "Pull from & Push to Coverity Connect — p.16-18",
        "Troubleshooting, Security, Checker/CWE Reference — p.19-21",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("⚠  Source code never leaves your machine. All analysis is local.")
    run.bold = True
    run.font.color.rgb = RGBColor(0xD9,0x77,0x06)
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFFBEB')
    pPr = p._p.get_or_add_pPr()
    pPr.append(shading)

    # Footer note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated for cross-team testing • For questions contact: Tool Owner (Coverity Findings Analyzer)")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x94,0xA3,0xB8)

    doc.add_page_break()

    # --- TABLE OF CONTENTS (manual) ---
    p = doc.add_heading("Table of Contents", level=1)
    p.runs[0].font.color.rgb = C_HDR
    toc_items = [
        ("1", "Overview — What the Tool Does & 5-Step Workflow", "3"),
        ("2", "System Requirements & Installation (EXE vs Python)", "3"),
        ("3", "Folder Structure — What to Ship to Other Teams", "4"),
        ("4", "Input Criteria — HTML Report / Excel / Server Pull (with expected formats)", "4"),
        ("5", "Source Code Root & Output Folder — Rules", "5"),
        ("6", "The Analysis Engine — What Happens Under the Hood", "6"),
        ("7", "GUI Tour 1: Setup Page (every field explained)", "7"),
        ("8", "GUI Tour 2: Analysis Page (progress, Elapsed/ETC, log)", "8"),
        ("9", "GUI Tour 3: Results Page (filters, tree, details, code viewer)", "9"),
        ("10", "Detail Window — Full Code View & Accept/Override", "10"),
        ("11", "Understanding Dispositions — Bug / False positive / Intentional / Needs review", "11"),
        ("12", "Output Files — CSV / JSON / Breakdown (exact columns)", "12"),
        ("13", "Pull Defects from Coverity Connect — Server → Excel", "13"),
        ("14", "Push Dispositions to Coverity — CSV → Server", "14"),
        ("15", "Decision Engine — Confidence, CWE/CERT/OWASP", "15"),
        ("16", "Troubleshooting & Performance Notes", "16"),
        ("17", "Security & Privacy", "16"),
        ("18", "Appendix A: Supported Checkers & CWE/CERT/OWASP Map", "17"),
        ("19", "Appendix B: Quick Start TL;DR", "18"),
        ("20", "Appendix C: EXE Build & Setup Instructions", "18"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Section"
    hdr[2].text = "Page"
    for c in hdr:
        for paragraph in c.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1E3A5F')
        c._tc.get_or_add_tcPr().append(shading)
    for num, title, pg in toc_items:
        row = table.add_row().cells
        row[0].text = num
        row[1].text = title
        row[2].text = pg
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph("Tip: Images in this guide have numbered red callouts. Match the numbers to the tables below each image.", style='Intense Quote').runs[0].font.size = Pt(8)

    # --- 1 Overview ---
    p = doc.add_heading("1. Overview — What the Tool Does", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("Coverity Findings Analyzer is a local desktop decision assistant for Coverity (STATIC) and Coverity Connect. It does not upload source code. It enriches each Coverity defect with the exact C/C++ function (tree-sitter), applies deterministic rules per checker (+ Z3/path-prover/flow-analysis where available), and proposes a triage: Bug / False positive / Intentional / Needs review, with confidence %, CWE/CERT/OWASP citation, and a code-exact Proposed Fix.")
    # 5-step workflow table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "Step"
    hdr[1].text = "Where"
    hdr[2].text = "Result"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'2563EB'); c._tc.get_or_add_tcPr().append(shading)
    steps = [
        ("1. Get defects", "HTML folder, Excel export, or ⬇ Pull from Coverity Connect", "Defect list: checker, file, line, type, severity"),
        ("2. Analyse", "Tree-sitter (once per file, cached) + rule engine per checker", "coverity_dispositions.csv (comment/fix/confidence/CWE)"),
        ("3. Review", "Results page → double-click row", "Events, source function, suggestion"),
        ("4. Decide", "✓ Accept Suggestion or ✎ Override", "coverity_final_decisions.csv"),
        ("5. Push back", "⬆ Push to Coverity (header)", "Server triage store updated"),
    ]
    for s,w,r in steps:
        row = table.add_row().cells
        row[0].text = s; row[1].text = w; row[2].text = r
        for c in row:
            for p in c.paragraphs:
                for run in p.runs: run.font.size=Pt(8)
    p = doc.add_paragraph()
    run = p.add_run("Output files (in your chosen Output Folder): "); run.bold=True; run.font.size=Pt(9)
    outs = [
        "coverity_dispositions.csv — machine suggestions (CID, Checker, Type, File, Line, Function, Classification, Comment, Fix, Timestamp, Category, CWE-url)",
        "audit.jsonl — full log (events, reasoning, confidence, context hash) — for audit",
        "coverity_final_decisions.csv — engineer-approved (Accepted/Overridden) — this is what Push uses",
        "needs_review_breakdown.txt — why Needs review (by reason/checker/category)",
        "coverity_pull_*.xlsx + _pull_log.txt — when you Pull from server",
    ]
    for o in outs:
        doc.add_paragraph(o, style='List Bullet').runs[0].font.size = Pt(8)

    # --- 2 System Requirements ---
    p = doc.add_heading("2. System Requirements & Installation", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_heading("2.1 For the Ready-Made EXE (recommended for testers)", level=2).runs[0].font.color.rgb = C_ACCENT
    bullets = [
        "Windows 10/11 64-bit, no Python needed",
        "Unzip Coverity-Tool-Setup.zip (see §3) → double-click CoverityTool.exe (or run_tool.bat)",
        "If Windows SmartScreen warns, click More info → Run anyway (exe is unsigned)",
        "First run may trigger antivirus scan — allow",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(9)
    doc.add_heading("2.2 For Python Run (developers)", level=2).runs[0].font.color.rgb = C_ACCENT
    doc.add_paragraph("Prerequisites: Python 3.10+ (python --version), C/C++ build tools only if a wheel fails (tree-sitter ships wheels). On Linux install tkinter: sudo apt install python3-tk").runs[0].font.size = Pt(9)
    # code block
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), 'F1F5F9'); pPr.append(shading)
    run = p.add_run("python -m venv .venv\n"
                    ".\\.venv\\Scripts\\Activate.ps1   # Windows\n"
                    "# source .venv/bin/activate      # Linux/macOS\n"
                    "pip install -r requirements.txt\n"
                    "python -c \"import bs4,lxml,openpyxl,zeep,tree_sitter_c; print('ok')\"\n"
                    "python local_gui.py   # MAIN GUI\n"
                    "# python coverity_triage.py --report <path> --src-root <path> --language cpp  # CLI")
    run.font.name = 'Consolas'; run.font.size = Pt(7.5); run.font.color.rgb = RGBColor(0x1E,0x29,0x3B)

    # --- 3 Folder Structure ---
    p = doc.add_heading("3. Folder Structure — What to Ship", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("Unzip and ship the entire setup folder. Do NOT move the exe out of its folder — it needs the _internal/ libs folder beside it (PyInstaller one-folder mode for faster start).").runs[0].font.size = Pt(9)
    # folder tree code block
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), 'F1F5F9'); pPr.append(shading)
    run = p.add_run(
        "Coverity-Tool-Setup/\n"
        "├── CoverityTool.exe          ← double-click here (or run_tool.bat)\n"
        "├── _internal/                ← Python libs, tree-sitter, z3 — do not delete\n"
        "├── run_tool.bat              ← sets Tcl/Tk paths and launches exe (use if exe fails)\n"
        "├── run_tool.ps1              ← PowerShell launcher\n"
        "├── docs/\n"
        "│   ├── Coverity_Tool_User_Guide.docx  ← this file\n"
        "│   └── images/               ← screenshots used in this guide\n"
        "├── sample_report/            ← optional: tiny HTML report for quick demo\n"
        "├── sample_src/               ← optional: 2-3 .c files to test source loading\n"
        "└── README.txt                ← 5-line quick start"
    )
    run.font.name = 'Consolas'; run.font.size = Pt(7.5)
    p = doc.add_paragraph()
    run = p.add_run("If you build the EXE yourself, see Appendix C: double-click build_exe.bat (needs PyInstaller) → dist/CoverityTool.exe"); run.font.size = Pt(8); run.italic = True

    # --- 4 Input Criteria ---
    p = doc.add_heading("4. Input Criteria — What Counts as Valid Input", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("The tool accepts exactly three kinds of input. It auto-detects which you gave. Nothing else is accepted.", style='Intense Quote').runs[0].font.size = Pt(9)

    # Input types table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, t in enumerate(["Input Type", "How to Select", "Expected Format", "What Happens If Wrong"]):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'1E3A5F'); hdr[i]._tc.get_or_add_tcPr().append(shading)
    rows = [
        ("HTML Report Folder", "Browse for FILE → pick any .html, or Browse for FOLDER → must contain index.html", "Coverity HTML export: folder with index.html + Code/*.html detail pages (or single .html). Parser tolerates any column order, finds CID/checker/file/line/type/severity by content.", "No table found in index.html → fix: open index.html, find table tag or provide sample"),
        ("Excel Export", "Browse → pick .xlsx / .xls", "Coverity/Black Duck Excel export. Required columns (fuzzy matched): CID/Defect ID, File, Checker/Type, plus optional Line, Function, Severity. Line = 'Various' → handled as function-scoped (see below).", "Could not find CID/Defect ID column → Headers: [...] — rename header to CID"),
        ("Pull from Coverity Connect (Server)", "Setup → ⬇ Pull from Coverity → fill Host/Port/User/Pass → Test Connection → Project → Stream → Pull", "Live server defect list via SOAP v9 + REST v2. Produces coverity_pull_<stream>_<stamp>.xlsx with integer Line (never Various) + pull log with _merged_line/_inst_line.", "SOAP Fault: Missing element sortAscending → fixed (v2025.9 pageSpec). If still fails, copy log's [shape,pageSize=] line"),
    ]
    for r in rows:
        row = table.add_row().cells
        for i, txt in enumerate(r):
            row[i].text = txt
            for p in row[i].paragraphs:
                for run in p.runs: run.font.size = Pt(7.5)
    # Various handling note
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), 'FFFBEB'); pPr.append(shading)
    run = p.add_run("About 'Various' line numbers (Excel only): "); run.bold=True; run.font.size=Pt(8)
    run = p.add_run("For function-scoped checkers (CHECKED_RETURN, UNUSED_VALUE, DEADCODE, MISSING_BREAK, INTEGER_OVERFLOW, ARRAY_VS_SINGLETON, STRING_NULL, etc.) the tool analyses the whole function and caps confidence. For memory-safety checkers (OVERRUN, BUFFER_SIZE, FORWARD_NULL) it asks you to provide the exact line rather than guess — those go to Needs review until you enter the line."); run.font.size=Pt(8)

    # Source code root rules
    p = doc.add_heading("5. Source Code Root & Output Folder — Rules", level=1)
    p.runs[0].font.color.rgb = C_HDR
    bullets = [
        "Source Code Root (required): folder containing your .c/.cpp/.h/.hpp files. The tool refuses to use the report folder itself as source root. If index.html is inside, it warns.",
        "It walks the tree once (cached) → extracts the exact function with tree-sitter (not regex). Files >500 KB and build folders (.git, build, out, node_modules, .venv) are skipped for speed.",
        "Output Folder: where coverity_dispositions.csv, audit.jsonl, coverity_final_decisions.csv, needs_review_breakdown.txt go. Defaults to Documents; each new run starts a fresh coverity_final_decisions.csv.",
        "Code Language: C++ (.cpp/.c) vs C only (.c) — picker on Setup page. Use C++ for mixed projects or run twice.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(8.5)
    # Highlight box
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), 'EFF6FF'); pPr.append(shading)
    run = p.add_run("✔ Tip: Point Source Root at the repo root (e.g., C:\\\\src\\\\myproj). The tool finds files by suffix-score, so Coverity absolute paths like /build/myproj/src/foo.c are resolved automatically."); run.font.size=Pt(8); run.italic=True

    # --- 6 Engine ---
    p = doc.add_heading("6. The Analysis Engine — Under the Hood (for curious testers)", level=1)
    p.runs[0].font.color.rgb = C_HDR
    bullets = [
        "1) parse_index_only() → light defect dicts (CID, checker, file, line, function, detail_file)",
        "2) warm_workspace_index() — one-time call-site index (scan once, not per defect) + tree-sitter parse cache by mtime",
        "3) For each defect: extract_enclosing_function() → build_defect_context() (callees/callers via workspace_indexer) → analyze_defect() dispatch per checker → DecisionAgent weighted evidence (Bug vs FP) + Z3/path_prover/flow_analysis where available → render_example_comment() → CWE/CERT/OWASP footer",
        "4) Confidence: dominance + margin + strongest-signal tie-break; caps for Various / fallback code",
        "5) Performance: indexing shows Elapsed/ETC separately; per-file semgrep cached & disabled by default (COVERITY_ENABLE_SEMGREP=1); concurrent REST discovery (8 workers, 3s) vs old 200s sequential",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(7.5)
    p = doc.add_paragraph("Security: verify_ssl=True by default (secure); checkbox Allow self-signed (insecure) for corporate certs. Passwords never logged; source never uploaded.", style='Intense Quote').runs[0].font.size = Pt(8)

    # Helper to add image with caption and highlight table
    def add_image_with_highlights(img_path, caption, highlights):
        if not os.path.exists(img_path):
            doc.add_paragraph(f"[Image missing: {img_path}]").runs[0].font.size = Pt(8)
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # Keep aspect, max width 6 inches
        run.add_picture(img_path, width=Inches(6))
        # Caption
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x64,0x74,0x8B)
        if highlights:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = "#"; hdr[1].text = "Highlight"; hdr[2].text = "What to Do"
            for c in hdr:
                for paragraph in c.paragraphs:
                    for r in paragraph.runs: r.bold=True; r.font.size=Pt(7); r.font.color.rgb=RGBColor(255,255,255)
                shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'2563EB'); c._tc.get_or_add_tcPr().append(shading)
            for num, h, do in highlights:
                row = table.add_row().cells
                row[0].text = str(num); row[1].text = h; row[2].text = do
                for c in row:
                    for paragraph in c.paragraphs:
                        for r in paragraph.runs: r.font.size = Pt(7)

    # --- 7 Setup Page ---
    p = doc.add_heading("7. GUI Tour 1: Setup Page — Every Field Explained", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("Launch: double-click CoverityTool.exe (or run_tool.bat) or python local_gui.py. You always start here.").runs[0].font.size = Pt(9)
    add_image_with_highlights(os.path.join(IMAGES_DIR, "gui-setup.png"),
        "Figure 1 — Setup Page. Inputs are validated before Start.",
        [
            ("1", "Coverity Report (HTML folder or Excel file)", "Browse for file (.html/.xlsx) or folder (must contain index.html). Blue ⬇ Pull from Coverity opens server pull dialog (§13)."),
            ("2", "Source Code Root (required)", "Browse to repo root with .c/.cpp/.h. Tool checks for ≥5 source files; warns if you picked report folder."),
            ("3", "Output Folder", "Where CSVs go. Defaults to Documents. Each new run clears coverity_final_decisions.csv."),
            ("4", "Code Language", "C++ for .cpp/.c mixed, C only for pure C. Affects tree-sitter parser."),
            ("5", "▶ Start Disposition", "Validates inputs → moves to Analysis page. Pre-flight errors show as popups + log."),
        ])
    # Validation table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "If You See..."; hdr[1].text = "Do This"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'DC2626'); c._tc.get_or_add_tcPr().append(shading)
    validations = [
        ("No index.html found in ...", "Pick the folder that contains index.html, not its parent. Check Coverity export completed."),
        ("Source Code Root cannot be same as input report", "Pick your repo root (with .c files), not the report."),
        ("Selected Source Root contains index.html", "You picked the report folder by mistake. Pick the source folder."),
        ("No C/C++ source files found", "Check the path; ensure .c/.cpp/.h exist. The tool will still run but all go to Needs review."),
    ]
    for a,b in validations:
        row = table.add_row().cells
        row[0].text = a; row[1].text = b
        for c in row:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(7.5)

    # --- 8 Analysis Page ---
    p = doc.add_heading("8. GUI Tour 2: Analysis Page — Progress, Elapsed/ETC, Log", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("After Start, the Analysis page runs in a background thread (UI stays responsive). Do not close the window until Done.").runs[0].font.size = Pt(9)
    add_image_with_highlights(os.path.join(IMAGES_DIR, "gui-analysis.png"),
        "Figure 2 — Analysis Page. Indexing is separate from per-defect ETA.",
        [
            ("1", "Reading index.html... → Found 42 defects", "First line: parse_index_only() quick scan. If 0 defects → Report is invalid."),
            ("2", "Progress bar + 0% (0/42 defects)", "Switches from indeterminate (indexing) to determinate (per-defect) after ready."),
            ("3", "Elapsed 0:00:12  ETC 0:00:45 (blue, top-right)", "Fixed: Elapsed includes indexing (not reset to 0). ETC is per-defect only (realistic). If >50ms/defect warning → source not loading."),
            ("4", "Log console (black)", "Live: Indexing source tree once (cached)... → Workspace indexed in 1.2s → ID 12345 [BUFFER_SIZE] Bug ... . Watch [Source] Loading: lines to see which file/checker is slow."),
            ("5", "Cancel", "Sets stop_evt; current defect finishes (semgrep up to 10s) then returns to Setup without saving."),
        ])
    # Performance notes
    p = doc.add_heading("Performance notes (what makes it fast now)", level=2)
    p.runs[0].font.color.rgb = C_ACCENT
    bullets = [
        "One-time workspace index (call-site + tree-sitter parse cache by mtime) — not per defect. Second run is fast.",
        "Skips build folders (.git, build, node_modules, .venv) and files >500 KB.",
        "Per-file semgrep cached & OFF by default (set COVERITY_ENABLE_SEMGREP=1 to enable). Previously 1000 defects ×10s = stuck.",
        "REST discovery concurrent (8 workers, 3s) vs old 40×5s = 200s stall.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(8)

    # --- 9 Results Page ---
    p = doc.add_heading("9. GUI Tour 3: Results Page — Filters, Tree, Details, Code Viewer", level=1)
    p.runs[0].font.color.rgb = C_HDR
    add_image_with_highlights(os.path.join(IMAGES_DIR, "gui-results.png"),
        "Figure 3 — Results Page. 3-column paned window: tree | details | code.",
        [
            ("1", "Top toolbar: Disposition Results + filter chips", "Click All/Bug/False positive/Intentional/Needs review/Accepted to filter. Category dropdown (Buffer overflow, Null pointer...) filters orthogonally. Count chips show totals."),
            ("2", "Left tree (ID + Class) grouped by Category", "Categories in CATEGORY_ORDER, Uncategorized last. Expand/collapse: click header. Double-click row → Detail Window. Tags: pushed=green, push_fail=red, conf_high/med/low."),
            ("3", "Middle details panel", "ID 12345 — BUFFER_SIZE (Buffer overflow) + File:line + Function | Severity | Confidence 85% (color by 0.8/0.6 thresholds) + Classification big colored + Comment paragraph (CWE/CERT footer) + Proposed Fix light-blue box (code-exact, // CWE tag) + buttons."),
            ("4", "Right code viewer (dark #1E1E1E)", "VS Code theme via Pygments. Line numbers gray, error line #5A1D1D red. Source origin banner: ✓ Local source file (green), ⚠ HTML-embedded (yellow), ✗ No source (red)."),
            ("5", "Bottom buttons: Full code view | Accept | Override", "Accept → status Accepted (green), writes coverity_final_decisions.csv. Override → dialog to pick classification + comment."),
        ])
    doc.add_paragraph("Tip: Use the filter drop-down to work the most important tail first (Bug → Needs review → False positive). The tree updates live.", style='Intense Quote').runs[0].font.size = Pt(8)

    # Table for detail panel fields
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i,t in enumerate(["Field in Middle Panel", "Meaning", "Action"]):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'1E3A5F'); hdr[i]._tc.get_or_add_tcPr().append(shading)
    fields = [
        ("ID 12345 — BUFFER_SIZE (Buffer overflow)", "CID + checker + category", "Unique ID for Coverity"),
        ("File: C:/.../foo.c : line 37 | Function: bar | Severity: High | Confidence: 85%", "Metadata + evidence strength", "Confidence ≥0.85 green, 0.6-0.8 orange, <0.6 red"),
        ("Classification: Bug (big colored)", "Machine suggestion: Bug / False positive / Intentional / Needs review → Accepted after you click", "Read color: Bug red #DC2626, FP green #16A34A, Intentional orange #D97706, Needs review blue #0891B2"),
        ("Comment paragraph", "Expert narrative: root cause → taint → guard → impact, with CWE-... (CERT ..., OWASP ...) footer", "Copy with Ctrl+C even when disabled; Select all Ctrl+A"),
        ("Proposed Fix (light-blue Consolas box)", "Code-exact fix with real variables, e.g., strncpy(buf,input,sizeof(buf)-1); buf[sizeof(buf)-1]='\\0'; // CWE-120", "Only for Bug/Needs review; Not shown for No fix required"),
        ("Source code viewer (right)", "Dark theme, error line red, origin banner", "Scroll to error line auto-centered; Copy works on disabled text"),
    ]
    for a,b,c in fields:
        row = table.add_row().cells
        row[0].text = a; row[1].text = b; row[2].text = c
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(7.5)

    # --- 10 Detail Window ---
    p = doc.add_heading("10. Detail Window — Full Code View & Accept/Override", level=1)
    p.runs[0].font.color.rgb = C_HDR
    add_image_with_highlights(os.path.join(IMAGES_DIR, "gui-detail.png"),
        "Figure 4 — Detail Window (double-click row). Full-screen, VS Code theme, same Accept/Override logic.",
        [
            ("1", "Top bar: ID 12345 BUFFER_SIZE Bug badge", "Color = classification; shows Accepted by user@time if already accepted"),
            ("2", "Left: Analysis + Proposed Fix + buttons", "Accept Suggestion → Accepted, Override → dialog (pick Bug/FP/Intentional/Needs review + comment)"),
            ("3", "Right: Source Code dark viewer", "Full function, not just snippet. Banner green/yellow/red for local/html/none. Line numbers + red error line + syntax. Tabs = 4 spaces."),
            ("4", "Copy", "Ctrl+C copies selected text even on disabled widgets; Ctrl+A selects all"),
        ])
    doc.add_paragraph("Flow: Review → Accept (writes to coverity_final_decisions.csv immediately, summary chips update) or Override → Save Override (choose classification, add comment) → Close → Next defect via tree or filter.", style='Intense Quote').runs[0].font.size = Pt(8)

    # --- 11 Understanding Dispositions ---
    p = doc.add_heading("11. Understanding Dispositions — What Each Means", level=1)
    p.runs[0].font.color.rgb = C_HDR
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i,t in enumerate(["Disposition", "Meaning", "When Tool Gives It", "Reviewer Action"]):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'2563EB'); hdr[i]._tc.get_or_add_tcPr().append(shading)
    rows = [
        ("Bug\n(red)", "Real defect at flagged location", "Evidence strongly for bug (dominance ≥0.62, margin ≥0.12, or critical label like strcpy without bounds, confirmed OOB)", "Fix it (see Proposed Fix) or confirm and Accept"),
        ("False positive\n(green)", "Rule inspected trace and concluded benign", "Guard dominates all paths (if(buf && idx<MAX)), or bounded API with sizeof-1 + explicit NUL, or caller bounds proven", "Accept (no fix)"),
        ("Intentional\n(orange)", "Code consciously does this", "(void) cast, ignore comment, #if 0, documented fallthrough, etc.", "Accept (no fix)"),
        ("Needs review\n(blue)", "Honest 'I don't know'", "No source/function, checker has no rule, or truly conflicting evidence, or line Various for memory-safety checker", "Manual review; add line number if Various"),
        ("Accepted\n(teal)", "You accepted suggestion", "After you click Accept", "Already recorded in coverity_final_decisions.csv"),
    ]
    for d,m,w,a in rows:
        row = table.add_row().cells
        row[0].text = d; row[1].text = m; row[2].text = w; row[3].text = a
        for c in row:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(7.5)
    # Confidence
    p = doc.add_paragraph()
    run = p.add_run("Confidence: "); run.bold=True; run.font.size=Pt(9)
    run = p.add_run("≥0.85 High (green) → auto-close after sanity check; 0.65-0.85 Moderate (orange) → quick walkthrough; <0.65 Low (red) → manual walkthrough. The phrase “High confidence” etc. is also the first sentence in the Comment.")
    run.font.size = Pt(8.5)

    # --- 12 Output Files ---
    p = doc.add_heading("12. Output Files — Where to Find What (in Output Folder)", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("Default Output Folder: Documents (you pick on Setup). All CSVs are UTF-8, overwrite per run except coverity_final_decisions.csv which dedupes by CID (last write wins).", style='Intense Quote').runs[0].font.size = Pt(8)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i,t in enumerate(["File", "Exact Columns / Format", "Used For"]):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'1E3A5F'); hdr[i]._tc.get_or_add_tcPr().append(shading)
    outs = [
        ("coverity_dispositions.csv", "CID, Checker, Type, Severity, File, Line, Function, Classification, Comment, Fix, Timestamp, Category, CWE-url (added in v7)", "Machine suggestions for every defect; import to Excel; feed to Push dialog"),
        ("audit.jsonl", "One JSON per line: {cid, checker, events, reasoning[], confidence, context_hash, code_start_line} — full evidence", "Audit trail, debugging, re-triage"),
        ("coverity_final_decisions.csv", "CID, Checker, File, Line, FinalClassification, FinalComment, Fix, Reviewer, Timestamp, Status (Accepted/Overridden), Category", "Engineer-approved; this is what Push pushes; deduped by CID"),
        ("needs_review_breakdown.txt", "Classification counts: Bug:12, ...\\nNeeds review rows: 8\\n  by reason: no_code:3, line_various:2...\\n  by checker: OVERUN:4...\\n  by category: Buffer overflow:5...", "Quick why-Needs-review summary without re-running"),
        ("coverity_pull_<stream>_<stamp>.xlsx", "Sheet Coverity: CID | Checker | Subtype | Severity | File | Line (integer, never Various) | Function | Events Summary", "Structured pull output — directly re-importable as Excel input; also shows pull correctness"),
        ("coverity_pull_*_pull_log.txt", "Pull log: Stream, Defects, Line numbers resolved, Events fetched, REST correction count, SOAP field dump (_merged_line/_inst_line/_main_event_line), per-CID final vs sources, 5 events each", "Post-pull diagnosis; attach when filing bug about wrong line numbers"),
    ]
    for f,fmt,use in outs:
        row = table.add_row().cells
        row[0].text = f; row[1].text = fmt; row[2].text = use
        for c in row:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(7)
    # Example CSV row
    p = doc.add_heading("Example coverity_dispositions.csv row (Excel view)", level=2)
    p.runs[0].font.color.rgb = C_ACCENT
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i,t in enumerate(["CID", "Checker", "Line", "Classification", "Confidence", "Comment (truncated)"]):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(7); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'2563EB'); hdr[i]._tc.get_or_add_tcPr().append(shading)
    row = table.add_row().cells
    row[0].text = "12345"; row[1].text = "BUFFER_SIZE"; row[2].text = "37"; row[3].text = "Bug"
    row[4].text = "85%"; row[5].text = "strncpy() size equals destination (64==64) — no room for NUL; buf is stack 64 bytes. Coverity trace: dest 64, copy 64. CWE-120..."
    for c in row:
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(7)

    # --- 13 Pull ---
    p = doc.add_heading("13. Pull Defects from Coverity Connect — Server → Excel (no 'Various')", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("Why Pull? Excel exports from the Coverity web UI often show Line = 'Various' (the defect spans multiple lines/macro). The server knows the current line (mainEventLineNumber). Pull fetches that via SOAP v9 + REST v2 and writes a clean Excel with integer lines.").runs[0].font.size = Pt(9)
    add_image_with_highlights(os.path.join(IMAGES_DIR, "gui-pull.png"),
        "Figure 5 — Pull Dialog (4 Sections). REST is tried first, SOAP is fallback.",
        [
            ("1", "Section 1 — Server Connection", "Host (e.g., coverity-er.honaero.com), Port 443, Username, Password → Test Connection. Toggle: Allow self-signed certificate (insecure — verify off) → checked for corporate self-signed, unchecked for production valid cert."),
            ("2", "Section 2 — Project & Stream", "Project dropdown (populated after Test Connection) → Stream dropdown (All streams in project + each stream). Defect limit spinner (100-50000, default 5000)."),
            ("3", "Section 3 — Output File", "Save path: coverity_pull_<stream>_<stamp>.xlsx (auto in Output Folder, Browse to change). Read-only, pre-filled."),
            ("4", "Section 4 — Pull Defects", "Check Fix current lines via Connect REST API (recommended — overlays mainEventLineNumber) → Test REST (probes /api/v2/streams with concurrent 8 workers) → ⬇ Pull Defects (progress bar + log). Log shows Line numbers resolved: 42/42 | Events fetched + REST current-line correction count."),
        ])
    # Steps
    doc.add_heading("Pull steps", level=2).runs[0].font.color.rgb = C_ACCENT
    steps = [
        "1. Setup → ⬇ Pull from Coverity → fill Host/Port/User/Pass → Test Connection (should show ✓ Connected (X projects found))",
        "2. Pick Project → Stream (pick single stream for fastest; All streams aggregates and dedupes by CID)",
        "3. Leave Fix current lines via REST checked (it corrects SOAP's stale lineNumber with REST's mainEventLineNumber)",
        "4. Click ⬇ Pull Defects → watch log: Trying REST API… → REST(POST /issues/search): 42 defects fetched → Fetching events… 42/42 → Writing Excel… → Saved : ...xlsx + Log: ...pull_log.txt",
        "5. If REST unavailable (log: REST unavailable (...), falling back to SOAP…), SOAP still works — lines come from defectInstance.lineNumber (last history instance) matching web UI",
        "6. On Done, dialog stays open so you can read log → Close → Setup's Coverity Report field is auto-filled with the generated Excel path → ▶ Start Disposition",
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number').runs[0].font.size = Pt(8)
    # REST vs SOAP note
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), 'EFF6FF'); pPr.append(shading)
    run = p.add_run("Note: REST discovery is now concurrent (8 workers, 3s timeout) vs old 40×5s sequential = 200s stall. The first Pull after Test Connection may still take 5-10s to probe /connect, /ngweb, 443/8443 ports — this is cached as _rest_base_discovered for the session.")
    run.font.size = Pt(7.5); run.italic = True

    # --- 14 Push ---
    p = doc.add_heading("14. Push Dispositions to Coverity — CSV → Server", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("Push writes your Accepted/Overridden decisions back to the Coverity triage store (Classification + Comment). It is available from the header ⬆ Push to Coverity at any time — you don't need to re-run analysis.").runs[0].font.size = Pt(9)
    add_image_with_highlights(os.path.join(IMAGES_DIR, "gui-push.png"),
        "Figure 6 — Push Dialog (3 Steps). Validates CIDs exist in target project.",
        [
            ("1", "Step 1 — Server Connection", "Same Host/Port/User/Pass + Allow self-signed toggle → Test Connection → Connected (X projects found)"),
            ("2", "Step 2 — Project & Stream + Triage Store", "Project dropdown → Stream dropdown → Triage Store (auto-filled as <project>-TS, editable; usually matches project name, not Default)"),
            ("3", "Step 3 — Load CSV & Review", "Browse… → pick coverity_dispositions.csv or coverity_final_decisions.csv (CID + Classification required) → table shows CID | ServerCID | Classification | Comment (trunc) | Checker | File — double-click row to edit Classification/Comment before pushing"),
            ("4", "Validate CIDs against Server", "🔍 Validate fetches all defects for the project and matches CSV CIDs by CID or (checker, basename). Green = matched, Red = NOT FOUND (removed — belongs to other project). Auto-validates after load."),
            ("5", "⬆ Push to Coverity", "Pushes in batches of 100 CIDs via updateTriageForCIDsInTriageStore → shows Succeeded/Failed + first error + tip: Try triage store = '<project>'"),
        ])
    steps = [
        "1. Header → ⬆ Push to Coverity → fill Host/Port/User/Pass → check Allow self-signed if needed → Test Connection",
        "2. Pick Project → Stream → verify Triage Store (if push fails, try store = project name, not Default; use ... button to list DefectService SOAP methods for debugging)",
        "3. Step 3 → Browse… → pick your CSV (the tool accepts both dispositions.csv and final_decisions.csv layouts, dedupes by CID keeping last)",
        "4. Wait for auto-validate (or click 🔍 Validate) → green rows = ready to push, red rows auto-removed",
        "5. Double-click any row to edit Classification/Comment if needed → then ⬆ Push to Coverity → watch Pushing 1/12… → Push Complete: Succeeded: 12, Failed: 0",
        "6. If Failed >0 and error is Triage Store, change store and retry; if CID NOT FOUND, you picked wrong project",
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number').runs[0].font.size = Pt(8)

    # --- 15 Decision Engine ---
    p = doc.add_heading("15. Decision Engine — Confidence, CWE/CERT/OWASP", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("The engine is not ML — it is weighted evidence with Z3/path-prover/flow-analysis where available. Every disposition cites CWE as primary taxonomy.").runs[0].font.size = Pt(9)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i,t in enumerate(["Classification", "What It Means", "Confidence & Evidence"]):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'2563EB'); hdr[i]._tc.get_or_add_tcPr().append(shading)
    rows = [
        ("Bug", "Real defect", "Confidence = 0.55+winner*0.22+dominance*0.15 etc. → capped. High ≥0.85 → close after sanity check."),
        ("False positive", "Guard dominates all paths, or bounded API with sizeof-1 + NUL, or caller bounds proven", "Critical FP labels (guard_dominates_all_paths, safe_bounded_api_with_sizeof) +0.15 if bug<0.3"),
        ("Intentional", "Deliberately ignored: (void) cast, ignore comment, #if 0, fallthrough annotation", "Only for CHECKED_RETURN, DEADCODE, NO_BREAK, etc. — not for BUFFER_SIZE"),
        ("Needs review", "Honest 'I don't know' — no source/function, or conflicting evidence, or line Various", "Honest fallback; reason shown in needs_review_breakdown.txt"),
    ]
    for a,b,c in rows:
        row = table.add_row().cells
        row[0].text = a; row[1].text = b; row[2].text = c
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(7.5)
    p = doc.add_paragraph()
    run = p.add_run("CWE example: BUFFER_SIZE → CWE-120 Buffer Copy without Checking Size (CERT STR31-C, A03:2021 Injection, CVSS 8.1, https://cwe.mitre.org/data/definitions/120.html). The Comment paragraph starts with CWE-120 and ends with Reference: CWE-120 | ... | URL. Proposed Fix is code-exact (e.g., strncpy(buf,src,sizeof(buf)-1); buf[sizeof(buf)-1]='\\0'; // CWE-120) — just suggestion, copy-pasteable with real variables.")
    run.font.size = Pt(8); run.italic = True

    # --- 16 Troubleshooting ---
    p = doc.add_heading("16. Troubleshooting & Performance", level=1)
    p.runs[0].font.color.rgb = C_HDR
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "Symptom"; hdr[1].text = "Fix"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'DC2626'); c._tc.get_or_add_tcPr().append(shading)
    issues = [
        ("Tool feels slow / stuck at 0 / N analysing...", "Indexing source tree once (cached) — wait for 'Workspace indexed in Xs — starting per-defect analysis.' ETA is per-defect only (indexing excluded) so it is realistic. Second run is fast. Per-file semgrep is OFF by default (COVERITY_ENABLE_SEMGREP=1 to enable)."),
        ("No table found in index.html", "HTML layout differs. Open index.html, find <table>, adjust html_report_parser.py or provide sample."),
        ("File not found for source files", "Coverity report paths must match on-disk layout. Set Source Code Root correctly; tool also searches by suffix-score and basename."),
        ("Tree-sitter fails to parse my C++ code", "Use C++ (.cpp/.c) language option. For mixed projects run twice."),
        ("Pull fails with SOAP Fault", "Copy log's [shape,pageSize=] line. Missing element sortAscending is now handled (all pageSpec shapes include sortAscending with exact casing)."),
        ("Pull REST 0 defects but SOAP has defects", "REST base discovery probed ports/roots concurrently (8 workers, 3s). If REST unavailable, it falls back to SOAP and overlays REST lines if Pull dialog's Fix current lines is checked."),
        ("Push: 0 Succeeded, Failed: N, First error: Triage Store", "Try store = project name (e.g., MyProject-TS), not Default. Use ... to list getTriageStores()."),
        ("Needs review for everything", "Intended for checkers with no rule yet or line Various. Add rule in heuristic_analyzer.py dispatch or provide concrete line."),
        ("Analysis completed very quickly (<50ms/defect)", "Source files not loading. Check Source Root path; watch log for [Source] Loading: lines."),
    ]
    for s,f in issues:
        row = table.add_row().cells
        row[0].text = s; row[1].text = f
        for c in row:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(7.5)

    doc.add_heading("Performance numbers (post-fix v7)", level=2).runs[0].font.color.rgb = C_ACCENT
    bullets = [
        "Workspace call-site index: built once, O(files) not O(defects). Before: per-defect full-tree walk → stuck on large reports. After: one walk, cached.",
        "Per-file tree-sitter parse cache keyed by mtime → defects in same file don't re-parse.",
        "Workspace symbol index (_INDEX_CACHE) cached per src_root+language.",
        "File content LRU (500 files, skip >1MB) → ~25 MB cap, no OOM.",
        "Skip build dirs (.git, build, node_modules, .venv) and >500 KB files.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(8)

    # --- 17 Security ---
    p = doc.add_heading("17. Security & Privacy", level=1)
    p.runs[0].font.color.rgb = C_HDR
    bullets = [
        "Source code never leaves your machine. Pull/Push only sends defects (CID, file, line, checker) and your triage decisions (Classification, Comment) to your Coverity Connect server.",
        "verify_ssl=True by default (secure). For corporate self-signed certs, check Allow self-signed certificate (insecure — verify off) in Pull/Push dialogs. Do not use insecure for production with valid cert.",
        "Passwords/tokens are held in memory only, never logged. The log shows _merged_line/_inst_line etc., not credentials.",
        "InsecureRequestWarning is suppressed only for the session with verify off, not globally.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(8)

    # --- Appendix A ---
    p = doc.add_heading("Appendix A: Supported Checkers & CWE/CERT/OWASP Map", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("Categories (checker_categories.py CATEGORY_ORDER): Buffer overflow, Memory - corruptions, Memory - illegal accesses, Null pointer dereferences, Integer handling, Resource leaks, Error handling, Control flow / code quality, Uncategorized (fallback).").runs[0].font.size = Pt(8)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i,t in enumerate(["Category", "Checker", "CWE", "CERT"]):
        hdr[i].text = t
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(7); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'1E3A5F'); hdr[i]._tc.get_or_add_tcPr().append(shading)
    # Load from cwe_mapping for accuracy
    try:
        import sys
        sys.path.insert(0, '/home/user/Coverity-Tool')
        from cwe_mapping import CWE_MAP
        from checker_categories import CHECKER_CATEGORIES
        for cat, chks in CHECKER_CATEGORIES.items():
            for chk in chks:
                info = CWE_MAP.get(chk, {})
                cwe = f"CWE-{info['cwe_id']} {info['cwe_name']}" if info else "—"
                cert = info.get('cert','—') if info else "—"
                row = table.add_row().cells
                row[0].text = cat; row[1].text = chk; row[2].text = cwe; row[3].text = cert
                for c in row:
                    for p in c.paragraphs:
                        for r in p.runs: r.font.size = Pt(6)
    except Exception as e:
        doc.add_paragraph(f"Could not load CWE map: {e}").runs[0].font.size = Pt(8)

    # --- Appendix B ---
    p = doc.add_heading("Appendix B: Quick Start TL;DR", level=1)
    p.runs[0].font.color.rgb = C_HDR
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), 'F1F5F9'); pPr.append(shading)
    run = p.add_run(
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1   # Windows\n"
        "pip install -r requirements.txt\n"
        "python local_gui.py\n"
        "1. Point Coverity Report at index.html folder (or ⬇ Pull from Coverity)\n"
        "2. Set Source Code Root to your C/C++ sources (repo root)\n"
        "3. ▶ Start Disposition → wait → double-click and Accept / Override each finding\n"
        "4. ⬆ Push to Coverity to write decisions back, or Export to Excel\n"
        "* Source code never leaves your machine *"
    )
    run.font.name = 'Consolas'; run.font.size = Pt(7.5)

    # --- Appendix C ---
    p = doc.add_heading("Appendix C: EXE Build & Setup Instructions", level=1)
    p.runs[0].font.color.rgb = C_HDR
    doc.add_paragraph("For teams without Python, ship the EXE one-folder build. Build on Windows (PyInstaller needs Windows to make .exe).").runs[0].font.size = Pt(9)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "File in Setup Folder"; hdr[1].text = "Purpose"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        shading = OxmlElement('w:shd'); shading.set(qn('w:fill'),'2563EB'); c._tc.get_or_add_tcPr().append(shading)
    files = [
        ("build_exe.bat", "Double-click to build: creates venv, pip install, pyinstaller CoverityTool.spec → dist/CoverityTool/CoverityTool.exe"),
        ("CoverityTool.spec", "PyInstaller spec (one-folder, --windowed, hidden imports: zeep, lxml, tree_sitter, z3, etc., datas: tcl/tk)"),
        ("run_tool.bat / run_tool.ps1", "Launchers that set TCL_LIBRARY/TK_LIBRARY and run exe; use if direct exe fails with _tkinter error"),
        ("requirements.txt", "Pinned versions (lxml==5.2.2, zeep==4.2.1, z3-solver==4.12.6.0, tree-sitter==0.23.2...)"),
        ("README.txt", "5-line quick start for testers (unzip → double-click)"),
    ]
    for f,desc in files:
        row = table.add_row().cells
        row[0].text = f; row[1].text = desc
        for c in row:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(7.5)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), 'FFFBEB'); pPr.append(shading)
    run = p.add_run("Build steps on Windows: 1) Install Python 3.10+ from python.org (check Add to PATH) 2) Unzip Coverity-Tool-Setup.zip 3) Double-click build_exe.bat → wait 2-4 min → dist/CoverityTool/CoverityTool.exe appears. Ship the whole dist/CoverityTool folder (exe + _internal) as a zip. Testers just unzip and double-click CoverityTool.exe — no Python needed.")
    run.font.size = Pt(8); run.bold = True

    # Add page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('w:space'),'preserve'); instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
        run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
        run = p.add_run(" | Coverity Findings Analyzer v1.4 — Confidential")
        run.font.size = Pt(7); run.font.color.rgb = RGBColor(0x94,0xA3,0xB8)

    # Save
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Saved to {OUT}, size {os.path.getsize(OUT)}")

if __name__ == "__main__":
    create_doc()
